# -*- coding: utf-8 -*-
"""批量导入模块：扫描导入源，自动归类、打标、生成规范笔记。

支持类型：
  - .md / .txt      → 内容成为笔记正文（txt 自动转 Markdown）
  - .pdf            → 默认提取正文生成笔记（pdf_mode=extract），原件归档附件；
                      也可仅归档（pdf_mode=archive）
  - .docx           → 尝试提取正文生成笔记，失败则归档附件
  - 图片等附件      → 直接归档「06附件/」

流程：扫描 → 哈希去重（幂等）→ 归类 → 打标 → 生成 frontmatter → 写笔记
      → 源文件移入「已处理」或「附件」→ 登记注册表 → 输出日志。
"""
import datetime
import logging
import os
import re
import shutil
from typing import Any, Dict, List, Optional, Tuple

from . import frontmatter, tagger as tagger_mod

CATEGORY_BY_EXT = {
    ".md": "Markdown 笔记",
    ".txt": "文本素材",
    ".pdf": "PDF 文档",
    ".docx": "Word 文档",
    ".xlsx": "表格文档",
    ".xls": "表格文档",
    ".pptx": "演示文稿",
    ".ppt": "演示文稿",
    ".png": "图片素材",
    ".jpg": "图片素材",
    ".jpeg": "图片素材",
    ".gif": "图片素材",
    ".webp": "图片素材",
}

ILLEGAL_CHARS = re.compile(r'[\\/:*?"<>|\x00-\x1f]')


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------
def _iter_files(root: str, max_depth: int = 5) -> List[str]:
    """递归收集文件（按路径排序，稳定输出）。跳过隐藏目录（.obsidian 等）。"""
    results: List[str] = []
    root = os.path.abspath(root)
    if not os.path.isdir(root):
        return results
    base_depth = root.count(os.sep)

    for dirpath, dirnames, filenames in os.walk(root):
        depth = dirpath.count(os.sep) - base_depth
        # 过滤隐藏目录与注册表/日志目录
        dirnames[:] = [d for d in dirnames
                       if not d.startswith(".") and d != "处理日志"]
        if depth >= max_depth:
            dirnames[:] = []
        for fn in sorted(filenames):
            if fn.startswith("."):
                continue
            results.append(os.path.join(dirpath, fn))
    return sorted(results)


def sanitize_name(stem: str, cfg: Dict[str, Any]) -> str:
    """文件名规范化：清理非法字符、替换空格、限长。"""
    naming = cfg.get("naming", {})
    name = ILLEGAL_CHARS.sub("_", stem).strip().strip(". ")
    if naming.get("sanitize", True):
        sp = naming.get("space_replacement", "_") or "_"
        name = re.sub(r"\s+", sp, name)
    max_len = int(naming.get("max_len", 80))
    if len(name) > max_len:
        name = name[:max_len].rstrip("_ .")
    return name or "未命名笔记"


def build_note_name(stem: str, cfg: Dict[str, Any], used: set) -> str:
    """生成笔记文件名（日期前缀 + 规范化名），冲突时追加序号。"""
    naming = cfg.get("naming", {})
    base = sanitize_name(stem, cfg)
    if naming.get("date_prefix", True):
        prefix = datetime.datetime.now().strftime(
            naming.get("date_prefix_format", "%Y-%m-%d")) + "_"
    else:
        prefix = ""
    candidate = prefix + base
    i = 1
    while candidate in used:
        i += 1
        candidate = "%s%s_%d" % (prefix, base, i)
    used.add(candidate)
    return candidate + ".md"


def _extract_pdf_text(path: str, max_pages: int = 50) -> Optional[str]:
    """提取 PDF 正文。优先 PyMuPDF，回退 pypdf。失败返回 None。"""
    try:
        import pymupdf as fitz  # PyMuPDF >= 1.24 新入口
    except Exception:
        try:
            import fitz  # 旧版入口（1.28 起弃用）
        except Exception:
            fitz = None
    if fitz is not None:
        try:
            doc = fitz.open(path)
            pages = min(doc.page_count, max_pages)
            parts = []
            for i in range(pages):
                parts.append(doc.load_page(i).get_text())
            doc.close()
            text = "\n\n".join(parts).strip()
            return text or None
        except Exception:
            pass
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        parts = []
        for i, page in enumerate(reader.pages[:max_pages]):
            parts.append(page.extract_text() or "")
        text = "\n\n".join(parts).strip()
        return text or None
    except Exception:
        return None


def _extract_docx_text(path: str) -> Optional[str]:
    """提取 Word 正文（python-docx，可选依赖）。"""
    try:
        import docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip() or None
    except Exception:
        return None


def _move_safe(src: str, dst_dir: str) -> str:
    """移动文件到目标目录，冲突时自动加时间戳后缀。返回目标路径。"""
    os.makedirs(dst_dir, exist_ok=True)
    dst = os.path.join(dst_dir, os.path.basename(src))
    if os.path.exists(dst):
        stem, ext = os.path.splitext(os.path.basename(src))
        dst = os.path.join(dst_dir, "%s_%s%s" % (
            stem, datetime.datetime.now().strftime("%H%M%S"), ext))
    shutil.move(src, dst)
    return dst


def _write_note_atomic(path: str, text: str) -> None:
    """原子写：先写临时文件再替换，避免半截文件。newline='\\n' 防双回车。"""
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8", newline="\n") as f:
        f.write(text)
    os.replace(tmp, path)


def _rel(path: str, root: str) -> str:
    """相对库根路径（正斜杠），用于 frontmatter.source 与日志。"""
    return os.path.relpath(path, root).replace("\\", "/")


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def run_import(cfg: Dict[str, Any], vault_root: str, registry: Any,
               logger: logging.Logger, report: Any,
               src: Optional[str] = None,
               dry_run: bool = False,
               no_move: bool = False,
               no_dedupe: bool = False) -> None:
    """执行批量导入。"""
    import_cfg = cfg["import"]
    structure = cfg["structure"]
    b_dir = os.path.join(vault_root, structure.get("B_知识提炼", "03知识提炼"))
    done_dir = os.path.join(vault_root, structure.get("已处理", "已处理"))
    attach_dir = os.path.join(vault_root, structure.get("附件", "附件"))
    inbox_rel = import_cfg.get("inbox", "未处理")
    inbox = os.path.abspath(os.path.join(vault_root, inbox_rel))
    extra_sources = import_cfg.get("extra_sources") or []
    # 导入源列表：显式 src 优先；否则 未处理（导入源；额外源默认无）
    sources = []
    if src:
        sources.append(os.path.abspath(src))
    else:
        sources.append(inbox)
        seen = {inbox}
        for s in extra_sources:
            sp = os.path.abspath(os.path.join(vault_root, s))
            if sp not in seen:
                seen.add(sp)
                sources.append(sp)

    # 确保目标目录存在（兼容自定义目录/未建库场景）
    for d in (b_dir, done_dir, attach_dir):
        os.makedirs(d, exist_ok=True)

    include_exts = {e.lower() for e in import_cfg.get("include_exts", [])}
    attach_exts = {e.lower() for e in import_cfg.get("attachment_exts", [])}
    archive_exts = {e.lower() for e in import_cfg.get("archive_exts", [])}
    pdf_mode = import_cfg.get("pdf_mode", "extract")
    pdf_max_pages = int(import_cfg.get("pdf_max_pages", 50))
    max_depth = int(import_cfg.get("max_depth", 5))

    tagger = tagger_mod.Tagger(cfg)
    date_format = cfg["frontmatter"].get("date_format", "%Y-%m-%d %H:%M")
    status_new = cfg["frontmatter"].get("status_new", "待整理")
    used_names: set = set()

    for src_path in sources:
        if not os.path.isdir(src_path):
            logger.info("[导入] 导入源不存在，跳过: %s", src_path)
            continue
        files = _iter_files(src_path, max_depth)
        logger.info("[导入] 扫描到 %d 个文件（来源：%s）", len(files), src_path)

        for fpath in files:
            report.scanned += 1
            ext = os.path.splitext(fpath)[1].lower()
            fname = os.path.basename(fpath)
            rel_src = _rel(fpath, vault_root)

            # 1. 支持类型检查
            if ext not in include_exts:
                logger.info("[导入] 跳过（不支持类型 %s）: %s", ext or "(无扩展名)", fname)
                report.skipped += 1
                report.add_detail("SKIP", "%s 不支持的类型 %s" % (fname, ext))
                continue

            # 2. 幂等：内容哈希去重
            content_hash = registry.hash_file(fpath)
            if (not no_dedupe and import_cfg.get("dedupe_by_hash", True)
                    and registry.contains(content_hash)):
                logger.info("[导入] 跳过（重复文件）: %s", fname)
                report.skipped += 1
                report.skipped_dups += 1
                report.add_detail("SKIP", "%s 内容重复（哈希 %s…）" % (fname, content_hash[:8]))
                continue

            # 3. 附件类 → 直接归档
            if ext in attach_exts:
                if dry_run:
                    logger.info("[DRY] 归档附件 → %s/ : %s", "附件", fname)
                    report.add_detail("DRY", "%s → 附件" % fname)
                    continue
                try:
                    dst = _move_safe(fpath, attach_dir)
                    registry.mark(content_hash, _rel(dst, vault_root), rel_src)
                    report.attachments += 1
                    report.add_detail("OK", "%s → 附件" % fname)
                    logger.info("[导入] 附件归档 → %s", _rel(dst, vault_root))
                except Exception as e:
                    report.failed += 1
                    report.errors.append("%s: %s" % (fname, e))
                    logger.error("[导入] 归档失败 %s: %s", fname, e)
                continue

            # 4. 文档类 → 提取/整理为笔记
            try:
                _handle_document(
                    cfg, fpath, fname, ext, rel_src, content_hash,
                    b_dir, done_dir, attach_dir, vault_root,
                    attach_exts, archive_exts,
                    pdf_mode, pdf_max_pages, tagger, date_format, status_new,
                    used_names, registry, logger, report,
                    dry_run=dry_run, no_move=no_move)
            except Exception as e:
                report.failed += 1
                report.errors.append("%s: %s" % (fname, e))
                logger.error("[导入] 处理失败 %s: %s", fname, e)

    registry.save()
    logger.info("[导入] 完成：导入 %d，新增笔记 %d，附件 %d，跳过 %d（重复 %d），失败 %d",
                report.imported, report.new_notes, report.attachments,
                report.skipped, report.skipped_dups, report.failed)


def _handle_document(cfg: Dict[str, Any], fpath: str, fname: str, ext: str,
                     rel_src: str, content_hash: str, b_dir: str, done_dir: str,
                     attach_dir: str, vault_root: str,
                     attach_exts: set, archive_exts: set,
                     pdf_mode: str, pdf_max_pages: int, tagger: Any,
                     date_format: str, status_new: str, used_names: set,
                     registry: Any, logger: logging.Logger, report: Any,
                     dry_run: bool, no_move: bool) -> None:
    """把单个文档处理成规范笔记。"""
    # 4.1 取正文
    body: Optional[str] = None
    if ext == ".md":
        body = frontmatter.read_text_auto(fpath)
        # 剥离已有 frontmatter，统一由程序重建（保留用户字段由 ensure_frontmatter 处理）
        _, body, _ = frontmatter.parse_frontmatter(body)
    elif ext == ".txt":
        body = frontmatter.read_text_auto(fpath)
    elif ext == ".pdf":
        if pdf_mode == "extract":
            body = _extract_pdf_text(fpath, pdf_max_pages)
            if body is None:
                logger.warning("[导入] PDF 文本提取失败，按归档处理: %s", fname)
        else:
            body = None
    elif ext == ".docx":
        body = _extract_docx_text(fpath)

    # 4.2 PDF/DOCX 提取失败或 archive 模式 → 仅归档
    if body is None and (ext in archive_exts or ext == ".pdf"):
        if dry_run:
            logger.info("[DRY] 归档文档 → %s/ : %s", "附件", fname)
            report.add_detail("DRY", "%s → 附件（无正文）" % fname)
            return
        dst = _move_safe(fpath, attach_dir)
        registry.mark(content_hash, _rel(dst, vault_root), rel_src)
        report.attachments += 1
        report.add_detail("OK", "%s → 附件" % fname)
        logger.info("[导入] 文档归档 → %s", _rel(dst, vault_root))
        return

    if body is None:
        body = ""  # 极端兜底

    # 4.3 生成标题 / 标签 / 分类
    stem = os.path.splitext(fname)[0]
    tags = tagger.tag_for(fname, body)
    category = CATEGORY_BY_EXT.get(ext, "其他文档")
    note_name = build_note_name(stem, cfg, used_names)
    note_path = os.path.join(b_dir, note_name)

    if dry_run:
        logger.info("[DRY] 生成笔记 → %s/%s （标签: %s）", "03知识提炼", note_name,
                    ", ".join("#" + t for t in tags) or "无")
        report.add_detail("DRY", "%s → 03知识提炼/%s 标签[%s]" % (
            fname, note_name, ",".join(tags)))
        return

    # 4.4 源文件归位（先移动，保证"移动→写笔记→登记"的提交顺序：
    #     任一步失败都不会留下"未登记却已处理"的重复导入隐患）
    if not no_move:
        if ext in archive_exts or ext == ".pdf":
            dst = _move_safe(fpath, attach_dir)
            src_record = _rel(dst, vault_root)
        else:
            dst = _move_safe(fpath, done_dir)
            src_record = _rel(dst, vault_root)
    else:
        src_record = rel_src

    # 4.5 写笔记（frontmatter + 正文）
    title = sanitize_name(stem, cfg)
    fm_text = frontmatter.build_frontmatter({
        "title": title,
        "tags": tags,
        "status": status_new,
        "category": category,
        # source 记录移动后的源素材路径（02已处理/xxx），供链接引擎关联
        "source": src_record,
    })
    note_text = fm_text + body.strip() + "\n"
    _write_note_atomic(note_path, note_text)
    report.imported += 1
    report.new_notes += 1
    for t in tags:
        report.tags_generated[t] += 1
    report.add_detail("OK", "%s → 03知识提炼/%s（标签: %s）" % (
        fname, note_name, ",".join("#" + t for t in tags) or "无"))
    logger.info("[导入] 生成笔记 %s （标签: %s）", note_path, ",".join(tags) or "无")

    # 4.6 登记注册表（幂等）
    registry.mark(content_hash, _rel(note_path, vault_root), src_record)
